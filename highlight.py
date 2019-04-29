from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import os, re, sys

def main():
    print("running...")
    path = "Documents"
    files = os.listdir(path)
    print(files)
    count = 0
    text_file = open("output-stats.txt", "w")
    #text_file.close()
    for file in files:

        #if file != '.DS_Store' and file.startswith('~$') == 'false':
        if file != '.DS_Store':
            count+=1;
            document = Document("Documents/" + file)
            print(count)
            text_file.write("%s" % count)
            text_file.write(" - ")

            #make word bank and initialize counts
            words = ['shall', 'required', 'must']
            counts = [0, 0, 0]

            # tables = document.tables
            # for table in tables:
            #     for row in table.rows:
            #         #print(row.cells)
            #         for cell in row.cells:
            #             #print(cell)
            #             for paragraph in cell.paragraphs:
            #                 print(paragraph.text)
            useless = 0
            #iterate through all paragraphs
            for paragraph in document.paragraphs:

                #iterate through wordbank
                for idx, word in enumerate(words):
                    if word.lower() in paragraph.text.lower():

                        #create runlist (all runs in paragraph)
                        runList = []
                        for run in paragraph.runs:
                            runList.append(run) 
                        #print(runList)

                        paragraph.text = '  '
                        for i, run in enumerate(runList):
                            #print(run.text)

                            #print(run.text + runList[i+1].text)
                            if (i+1) <= len(runList)-1:
                                if not (word in run.text or word in runList[i+1].text):
                                    if  word in (run.text + runList[i+1].text):
                                        print("run: " + run.text)
                                        run.text = run.text + runList[i+1].text
                                        runList[i+1].text = ''
                                        print("run + 1: " + runList[i+1].text)
                                        print("new run: " + run.text)

                                        

                            if word.lower() in run.text.lower():
                                #split_runs = run.text.split(word)

                                split_runs = re.split("\\b(" + word + ")\\b(?i)", run.text)
                                #print(split_runs)
                                for i in range(len(split_runs)-1):
                                    if word.lower() in split_runs[i].lower():
                                        newRun = add_run_copy(paragraph, run, split_runs[i])
                                        font = newRun.font
                                        if (idx < 1):
                                            font.highlight_color = WD_COLOR_INDEX.YELLOW
                                            counts[0]+=1
                                        if (idx == 1):
                                            font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                                            counts[1]+=1                                            
                                        if (idx > 1):
                                            font.highlight_color = WD_COLOR_INDEX.PINK
                                            counts[2]+=1                                            
                                    else:
                                        add_run_copy(paragraph, run, split_runs[i])
                                    
                                    #newRun = add_run_copy(paragraph, run, word)
                                    #font = newRun.font
                                    
                                    #newRun.highlight_color = WD_COLOR_INDEX.YELLOW
                                    #newRun = paragraph.add_run('shall').font #add color
                                add_run_copy(paragraph, run, split_runs[len(split_runs)-1])
                            else: 
                                add_run_copy(paragraph, run)

            #iterate through all tables
            for table in document.tables:
                #print("BUZZ 1")
                for row in table.rows:
                    #print("BUZZ 2")
                    #print(row.cells)
                    for cell in row.cells:

                        for paragraph in cell.paragraphs:

                            #print(paragraph.text)
                            #iterate through wordbank
                            for idx, word in enumerate(words):
                                if word.lower() in paragraph.text.lower():

                                    #create runlist (all runs in paragraph)
                                    runList = []
                                    for run in paragraph.runs:
                                        runList.append(run) 
                                    #print(runList)

                                    paragraph.text = '  '
                                    for run in runList:
                                        #print(run.text)

                                        #print(run.text + runList[i+1].text)
                                        if (i+1) <= len(runList)-1:
                                            if not (word in run.text or word in runList[i+1].text):
                                                if  word in (run.text + runList[i+1].text):
                                                    print("run: " + run.text)
                                                    run.text = run.text + runList[i+1].text
                                                    runList[i+1].text = ''
                                                    print("run + 1: " + runList[i+1].text)
                                                    print("new run: " + run.text)


                                        if word.lower() in run.text.lower():
                                            #split_runs = run.text.split(word)

                                            split_runs = re.split("\\b(" + word + ")\\b(?i)", run.text)
                                            #print(split_runs)
                                            for i in range(len(split_runs)-1):
                                                if word.lower() in split_runs[i].lower():
                                                    newRun = add_run_copy(paragraph, run, split_runs[i])
                                                    font = newRun.font
                                                    if (idx < 1):
                                                        font.highlight_color = WD_COLOR_INDEX.YELLOW
                                                        counts[0]+=1
                                                    if (idx == 1):
                                                        font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                                                        counts[1]+=1                                            
                                                    if (idx > 1):
                                                        font.highlight_color = WD_COLOR_INDEX.PINK
                                                        counts[2]+=1                                            
                                                else:
                                                    add_run_copy(paragraph, run, split_runs[i])
                                                
                                                #newRun = add_run_copy(paragraph, run, word)
                                                #font = newRun.font
                                                
                                                #newRun.highlight_color = WD_COLOR_INDEX.YELLOW
                                                #newRun = paragraph.add_run('shall').font #add color
                                            add_run_copy(paragraph, run, split_runs[len(split_runs)-1])
                                        else: 
                                            add_run_copy(paragraph, run)


            #Save the main document
            document.save('./Output/' + file) 

            #write word count to output-stats 
            text_file.write("[ shall: %s, must: %s, required: %s ]" % (counts[0], counts[2], counts[1]) )
            text_file.write(" - ")
            text_file.write(file + "\r\n")

    text_file.close() #close the text file



def add_run_copy(paragraph, run, text=None):
    r = paragraph.add_run(text=run.text if text is None else text, style=run.style)
    r.bold = run.bold
    r.italic = run.italic
    r.underline = run.underline
    r.font.all_caps = run.font.all_caps
    r.font.bold = run.font.bold
    r.font.color.rgb = run.font.color.rgb
    r.font.color.theme_color = run.font.color.theme_color
    #r.font.color.type = run.font.color.type
    r.font.complex_script = run.font.complex_script
    r.font.cs_bold = run.font.cs_bold
    r.font.cs_italic = run.font.cs_italic
    r.font.double_strike = run.font.double_strike
    r.font.emboss = run.font.emboss
    r.font.hidden = run.font.hidden
    r.font.highlight_color = run.font.highlight_color
    r.font.imprint = run.font.imprint
    r.font.italic = run.font.italic
    r.font.math = run.font.math
    r.font.name = run.font.name
    r.font.no_proof = run.font.no_proof
    r.font.outline = run.font.outline
    r.font.rtl = run.font.rtl
    r.font.shadow = run.font.shadow
    r.font.size = run.font.size
    r.font.small_caps = run.font.small_caps
    r.font.snap_to_grid = run.font.snap_to_grid
    r.font.spec_vanish = run.font.spec_vanish
    r.font.strike = run.font.strike
    r.font.subscript = run.font.subscript
    r.font.superscript = run.font.superscript
    r.font.underline = run.font.underline
    r.font.web_hidden = run.font.web_hidden
    return r


main()